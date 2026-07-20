from __future__ import annotations

import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
EVIDENCE_PATH = REPO_ROOT / "docs" / "demo-assets" / "multi-household-validation-results.json"
SCREENSHOT_PATH = REPO_ROOT / "docs" / "demo-assets" / "multi-household-dashboard.png"
OVERVIEW_IMAGE_PATH = REPO_ROOT / "docs" / "demo-assets" / "multi-household-dashboard-overview.png"
SEMANTIC_CHART_PATH = REPO_ROOT / "docs" / "demo-assets" / "multi-household-semantic-outcomes.png"
OUTPUT_PATH = REPO_ROOT / "docs" / "multi-household-pipeline-validation-report.docx"

BLUE = "2E74B5"
DARK_BLUE = "17365D"
GREEN = "008A6A"
AMBER = "C47A16"
DARK = "17212B"
MID = "5C6773"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "E8F5F0"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, keep=True) -> None:
    paragraph.paragraph_format.keep_with_next = keep


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 28, DARK_BLUE, 0, 12),
        ("Heading 1", 16, BLUE, 16, 7),
        ("Heading 2", 13, DARK_BLUE, 12, 5),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.86))
    header_table.columns[0].width = Inches(4.8)
    header_table.columns[1].width = Inches(2.06)
    left = header_table.cell(0, 0).paragraphs[0]
    run = left.add_run("AD-FLEX  |  MULTI-HOUSEHOLD VALIDATION")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID)
    right = header_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right.add_run("20 JULY 2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID)
    for cell in header_table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": 8, "color": "D9DEE5"})

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Local validation environment  |  No real device execution  |  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID)
    add_field(paragraph, "PAGE")


def add_label(document: Document, text: str, color=GREEN) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(document: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)


def add_bullets(document: Document, items) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def add_callout(document: Document, title: str, text: str, fill=PALE_GREEN, accent=GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.72)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    set_cell_border(cell, left={"val": "single", "sz": 22, "color": accent})
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph.add_run(f"\n{text}")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(document: Document, headers, rows, widths=None, font_size=9) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = 1
    if widths:
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(header))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        set_cell_shading(row.cells[0], WHITE)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            set_cell_margins(cell)
            if row_index % 2 == 1:
                set_cell_shading(cell, LIGHT)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            run.font.size = Pt(font_size)
            if index > 0 and isinstance(value, (int, float)):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_kpi_grid(document: Document, metrics) -> None:
    table = document.add_table(rows=2, cols=4)
    table.autofit = False
    for index in range(4):
        table.columns[index].width = Inches(1.68)
    for row_index in range(2):
        for column_index in range(4):
            metric_index = row_index * 4 + column_index
            cell = table.cell(row_index, column_index)
            set_cell_margins(cell, top=130, start=130, bottom=130, end=130)
            set_cell_shading(cell, PALE_BLUE if row_index == 0 else LIGHT)
            set_cell_border(cell, bottom={"val": "single", "sz": 5, "color": "D9DEE5"})
            value, label = metrics[metric_index]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            run.bold = True
            run.font.size = Pt(17)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            label_paragraph = cell.add_paragraph()
            label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_paragraph.paragraph_format.space_after = Pt(0)
            run = label_paragraph.add_run(label)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MID)


def make_semantic_chart(evidence) -> None:
    labels = ["SLM primary", "Deterministic fallback", "Safely unmapped"]
    values = [
        evidence["pipeline"]["slm_primary"],
        evidence["pipeline"]["deterministic_fallback"],
        evidence["pipeline"]["unmapped"],
    ]
    colors = ["#2E74B5", "#008A6A", "#9AA4AF"]
    canvas = Image.new("RGB", (1500, 470), "white")
    draw = ImageDraw.Draw(canvas)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    label_font = ImageFont.truetype(str(font_path), 28) if font_path.exists() else ImageFont.load_default()
    value_font = ImageFont.truetype(str(bold_path), 28) if bold_path.exists() else label_font
    tick_font = ImageFont.truetype(str(font_path), 21) if font_path.exists() else ImageFont.load_default()
    chart_left = 410
    chart_right = 1400
    chart_width = chart_right - chart_left
    maximum = max(values) or 1
    for tick in range(0, maximum + 1, 20):
        x = chart_left + int(chart_width * tick / maximum)
        draw.line((x, 45, x, 395), fill="#E8EBEF", width=2)
        draw.text((x - 10, 410), str(tick), fill="#687481", font=tick_font)
    for index, (label, value, color) in enumerate(zip(labels, values, colors)):
        y = 70 + index * 112
        draw.text((18, y + 16), label, fill="#263442", font=label_font)
        bar_width = int(chart_width * value / maximum)
        draw.rounded_rectangle((chart_left, y, chart_left + bar_width, y + 64), radius=8, fill=color)
        draw.text((chart_left + bar_width + 16, y + 15), str(value), fill="#263442", font=value_font)
    canvas.save(SEMANTIC_CHART_PATH, optimize=True)


def make_overview_crop() -> None:
    if not SCREENSHOT_PATH.exists():
        return
    with Image.open(SCREENSHOT_PATH) as image:
        crop_height = min(image.height, 1050)
        crop = image.crop((0, 0, image.width, crop_height))
        crop.save(OVERVIEW_IMAGE_PATH, optimize=True)


def build_report() -> None:
    evidence = json.loads(EVIDENCE_PATH.read_text(encoding="utf-8"))
    make_semantic_chart(evidence)
    make_overview_crop()

    pipeline = evidence["pipeline"]
    generation = evidence["generation"]
    load = evidence["load_management"]
    configuration = evidence["configuration"]

    document = Document()
    configure_document(document)
    document.core_properties.title = "AD-FLEX Multi-Household Pipeline Validation Report"
    document.core_properties.subject = "Five-household local scalability validation"
    document.core_properties.author = "AD-FLEX Engineering"

    add_label(document, "Local Platform Validation")
    title = document.add_paragraph(style="Title")
    title.add_run("Multi-Household Pipeline\nValidation Report")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Five households | Fifteen simulated energy assets | End-to-end evidence")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(MID)

    add_kpi_grid(
        document,
        [
            (configuration["households"], "Households"),
            (configuration["devices"], "Devices"),
            (pipeline["raw_telemetry"], "Telemetry messages"),
            (pipeline["semantic_events"], "Semantic readings"),
            (f"{pipeline['overall_success_rate_percent']}%", "Completion rate"),
            (pipeline["slm_primary"], "SLM-primary mappings"),
            (pipeline["deterministic_fallback"], "Fallback mappings"),
            (pipeline["dropped_readings"], "Dropped readings"),
        ],
    )
    document.add_paragraph()
    add_callout(
        document,
        "Validated outcome",
        "The final two-update-per-device run completed with 120 unique semantic and IEEE records, zero drops, zero duplicate rows, zero processing errors, and no real device execution.",
    )
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.add_run(f"Validation run: {evidence['run_id']}\n").bold = True
    paragraph.add_run("Environment: Local Docker Compose, Kafka, TimescaleDB, Ollama Phi-3 Mini, and Next.js customer console\n")
    paragraph.add_run("Prepared: 20 July 2026")
    document.add_page_break()

    document.add_heading("1. Executive Summary", level=1)
    add_body(
        document,
        "This validation assessed whether the AD-FLEX local architecture could support multiple households at the same time without changing the existing business workflow. Five households were created, each with a Shelly Plug, an Enode / Easee EV charger, and a heat pump. Every simulated asset had a unique device and household identity, an independent update interval, changing values, and proper timestamps.",
    )
    add_body(
        document,
        "The final run submitted 30 telemetry messages through the security gateway and produced 120 normalized readings. Every reading reached semantic storage and IEEE 2030.5-style storage. No message or reading was dropped, no duplicate semantic or IEEE row was recorded, and no processing error occurred. The end-to-end completion rate was 100%.",
    )
    add_body(
        document,
        "Local Phi-3 Mini was called for all 120 readings. Eighty mappings passed as SLM-primary, ten known readings used deterministic fallback, and thirty were safely marked unmapped after failing configured guardrails. These rejected outputs remained visible and auditable without stopping the pipeline.",
    )
    add_callout(
        document,
        "Business interpretation",
        "The platform handled the intended production-style demonstration reliably. The main scale-up constraint is local serial model inference, not gateway acceptance, Kafka delivery, or database insertion.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    document.add_heading("2. Test Configuration", level=1)
    add_table(
        document,
        ["Configuration item", "Measured setup"],
        [
            ("Households", "5"),
            ("Energy assets", "15 total; 3 per household"),
            ("Device types", "5 Shelly Plugs, 5 Enode / Easee EV chargers, 5 heat pumps"),
            ("Updates", "2 per device"),
            ("Update interval range", f"{configuration['randomized_interval_ms']['minimum']}-{configuration['randomized_interval_ms']['maximum']} ms"),
            ("Ingress", "Security gateway POST /telemetry with local API key"),
            ("Semantic model", "Local Ollama phi3:mini, SLM primary"),
            ("Validation", "Deterministic SAREF4ENER guardrails and fallback"),
            ("Event and data services", "Kafka, PostgreSQL/TimescaleDB, MQTT"),
            ("Control mode", "Approval plus mock dispatch; no real device execution"),
            ("Customer application", "Next.js console at http://localhost:3000"),
        ],
        widths=[2.1, 4.6],
    )
    add_body(
        document,
        "All 15 devices produced two distinct telemetry samples and two ordered, unique timestamps. Shelly values covered active power, voltage, current, and cumulative energy. EV charger values covered charging state, charging power, and delivered energy. Heat-pump values covered room temperature, target temperature, operating mode, flow temperature, and power usage.",
    )

    document.add_heading("3. Validation Procedure", level=1)
    add_table(
        document,
        ["Step", "Procedure", "Evidence"],
        [
            (1, "Build and start all Docker Compose services", "18 containers running; zero restarts"),
            (2, "Check service, Kafka, TimescaleDB, and Ollama readiness", "Health checks passed; phi3:mini listed"),
            (3, "Create five households and fifteen independently seeded devices", "Unique IDs and 322-1,382 ms intervals"),
            (4, "Send two telemetry updates per device through the gateway", "30 of 30 requests accepted"),
            (5, "Wait for semantic and IEEE storage", "120 unique rows at each stage"),
            (6, "Run DSO signal and approval workflow", "Proposal 89 reached ready_to_dispatch"),
            (7, "Verify mock dispatch and device API translation", "1 mock row and 3 device command rows"),
            (8, "Run dataspace export", "Minimized, pseudonymized 100-record summary"),
            (9, "Publish a separate MQTT message", "1 raw and 4 rows at each reading stage"),
            (10, "Open the customer console with Playwright", "15 final-run devices and Phi-3 primary status visible"),
        ],
        widths=[0.55, 3.45, 2.7],
        font_size=8.5,
    )

    document.add_page_break()
    document.add_heading("4. Runtime Results", level=1)
    document.add_heading("Delivery and persistence", level=2)
    add_table(
        document,
        ["Metric", "Result", "Assessment"],
        [
            ("Telemetry generated / accepted", "30 / 30", "Pass"),
            ("Raw telemetry rows", pipeline["raw_telemetry"], "Pass"),
            ("Normalized readings", pipeline["normalized_telemetry"], "Pass"),
            ("Semantic rows / unique", f"{pipeline['semantic_events']} / {pipeline['semantic_unique_readings']}", "Pass"),
            ("IEEE rows / unique", f"{pipeline['ieee20305_events']} / {pipeline['ieee20305_unique_readings']}", "Pass"),
            ("Duplicate semantic / IEEE rows", f"{pipeline['semantic_duplicate_rows']} / {pipeline['ieee20305_duplicate_rows']}", "Pass"),
            ("Dropped messages / readings", f"{pipeline['dropped_messages']} / {pipeline['dropped_readings']}", "Pass"),
            ("Processing errors", pipeline["processing_errors"], "Pass"),
            ("Overall completion rate", f"{pipeline['overall_success_rate_percent']}%", "Pass"),
        ],
        widths=[3.0, 1.7, 1.3],
    )

    document.add_heading("Latency and throughput", level=2)
    add_table(
        document,
        ["Metric", "Minimum", "Average", "Maximum"],
        [
            ("Gateway response", "10.56 ms", "13.04 ms", "17.30 ms"),
            ("Raw database insert", "7.08 ms", "10.02 ms", "18.92 ms"),
            ("Semantic queue-to-persist", "2.22 s", "112.86 s", "222.26 s"),
            ("Raw-to-IEEE completion", "2.24 s", "112.87 s", "222.26 s"),
        ],
        widths=[2.6, 1.3, 1.3, 1.3],
    )
    add_bullets(
        document,
        [
            "Gateway input burst: approximately 7.143 telemetry messages per second.",
            "Observed end-to-end telemetry completion: 0.133 messages per second.",
            "Observed semantic/IEEE reading completion: 0.531 readings per second.",
            "Measured workflow uptime: 233.4 seconds, including DSO, dispatch, dataspace, and evidence collection.",
            "Semantic latency includes queue waiting behind serial local model inference; it is not pure model execution time.",
        ],
    )

    document.add_heading("Semantic mapping outcomes", level=2)
    document.add_picture(str(SEMANTIC_CHART_PATH), width=Inches(6.55))
    caption = document.add_paragraph("Figure 1. Final-run semantic outcomes across 120 readings.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(8.5)
    add_table(
        document,
        ["Source", "Confidence", "Count", "Share"],
        [
            ("SLM primary", "High", 53, "44.2%"),
            ("SLM primary", "Medium", 27, "22.5%"),
            ("Deterministic fallback", "High", 10, "8.3%"),
            ("Safely unmapped", "Low", 30, "25.0%"),
        ],
        widths=[2.4, 1.4, 1.1, 1.1],
    )

    document.add_heading("Database and workflow evidence", level=2)
    add_table(
        document,
        ["Evidence area", "Measured result"],
        [
            ("Security gateway audit", "35 run-specific rows"),
            ("Dispatch command", "1 row; proposal 89"),
            ("Approval audit", "3 rows"),
            ("Mock dispatch audit", "1 row; no_real_execution=true"),
            ("Device command audit", "3 rows; no_real_execution=true"),
            ("Dataspace export audit", "1 row in validation window"),
            ("Dataspace response", "100 records; minimized and pseudonymized"),
        ],
        widths=[2.6, 4.0],
    )

    document.add_heading("Resource observations", level=2)
    add_table(
        document,
        ["Component", "CPU snapshot", "Memory snapshot"],
        [
            ("Kafka", "0.86%", "571 MiB"),
            ("TimescaleDB", "0.02%", "99.29 MiB"),
            ("Zookeeper", "0.22%", "176.8 MiB"),
            ("Semantic connector", "2.24%", "37.75 MiB"),
            ("Security gateway", "2.04%", "29.93 MiB"),
            ("Highest sampled application CPU", "2.35%", "Dataspace export"),
        ],
        widths=[3.0, 1.7, 1.9],
    )
    add_body(document, "Resource values are point-in-time Docker samples and should not be interpreted as peak utilization.")

    document.add_page_break()
    document.add_heading("5. Pipeline Observations", level=1)
    observations = [
        ("Security Gateway", "Accepted every authenticated telemetry request. Average response time was 13.04 ms. Correlation IDs and safe request metadata were audited."),
        ("Kafka Digital Spine", "Delivered all telemetry and downstream events. The final run had no coordinator error, offset replay, or detected drop."),
        ("Processing Engine", "Expanded 30 device payloads into exactly 120 normalized readings."),
        ("Semantic Connector", "Called local Phi-3 Mini for every reading, applied deterministic validation, and stored all accepted, fallback, and unmapped outcomes."),
        ("Phi-3 Mini", "Was available locally and produced 80 guardrail-approved SLM-primary mappings. Local serial inference was the principal latency contributor."),
        ("TimescaleDB", "Stored all run-specific telemetry, semantic, IEEE, approval, dispatch, device-command, gateway, and export evidence. Historical queries returned all five households and fifteen devices."),
        ("IEEE 2030.5 Translator", "Created 120 simplified IEEE 2030.5-style rows. The implementation remains a translator foundation and is not certified."),
        ("Approval Workflow", "Enforced proposed to reviewed to approved to ready-to-dispatch transitions."),
        ("Mock Dispatch", "Created one simulated result with no_real_execution=true."),
        ("Device Command Translator", "Created three simulated device-specific results for Shelly, Enode / Easee, and heat pump adapters."),
        ("Customer Dashboard", "Displayed live SLM status, final-run device IDs, storage, security, dispatch, and dataspace information. Desktop and mobile checks passed."),
        ("Dataspace Export", "Returned a 100-record minimized and pseudonymized summary without raw private payloads."),
    ]
    for name, text in observations:
        document.add_heading(name, level=3)
        add_body(document, text)

    if OVERVIEW_IMAGE_PATH.exists():
        document.add_heading("Customer Console Evidence", level=2)
        document.add_picture(str(OVERVIEW_IMAGE_PATH), width=Inches(6.55))
        caption = document.add_paragraph("Figure 2. Live customer console overview after the multi-household run.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(8.5)

    document.add_heading("6. Scalability Assessment", level=1)
    add_body(
        document,
        "The local architecture successfully handled the intended five-household workload. Concurrent gateway intake, Kafka delivery, normalization, persistence, protocol translation, approval, mock dispatch, device API translation, dataspace export, and dashboard visibility all remained operational. The final run completed with zero drops, zero duplicate semantic/IEEE rows, and zero processing errors.",
    )
    add_body(
        document,
        "The main bottleneck is local serial semantic inference. The gateway accepted the input burst quickly, raw inserts averaged 10.02 ms, and application container resource samples remained modest. Queue-to-semantic latency increased because 120 readings waited behind a single local model consumer.",
    )
    add_callout(
        document,
        "Recommended next scale step",
        "Introduce bounded semantic concurrency, model warm-up, partition-aware consumers, per-inference latency metrics, and expanded deterministic coverage for device state codes before increasing the household count.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    document.add_heading("Finding Resolved During Validation", level=2)
    add_body(
        document,
        "An initial diagnostic run produced one duplicate semantic/IEEE row when a 30-second local model timeout exceeded KafkaJS's default consumer session and triggered offset redelivery. The semantic consumer now uses a configurable 120-second session and enforces a session at least 30 seconds longer than the SLM timeout. The final 120-reading run produced no duplicates and no coordinator errors.",
    )

    document.add_heading("7. Limitations", level=1)
    add_table(
        document,
        ["Current implemented local platform", "Future production deployment"],
        [
            ("Five-household, fifteen-device controlled demonstration", "Progressive tests with larger household populations and sustained durations"),
            ("Single local Phi-3 Mini instance", "Horizontally scalable inference workers, warm pools, and model observability"),
            ("Local API key and Docker networking", "Production identity, mTLS, secrets management, private networking, and WAF"),
            ("Simplified IEEE 2030.5-style translator", "Formal interoperability testing and certification where required"),
            ("IDS/ENERSHARE-ready export foundation", "Real connector, identity, contract negotiation, and external publication"),
            ("Mock dispatch and simulated device APIs", "Real device credentials, consent, operator controls, safety certification, and rollback"),
            ("Point-in-time resource samples", "Continuous CPU, memory, tracing, queue, and SLO monitoring"),
            ("Basic ontology coverage for state codes", "Expanded deterministic/SAREF4ENER mappings and governed vocabulary"),
        ],
        widths=[3.25, 3.35],
        font_size=8.4,
    )

    document.add_heading("8. Conclusions", level=1)
    add_body(
        document,
        "The AD-FLEX local pipeline successfully supported approximately fifteen simultaneously represented energy assets across five households. Two independent updates from every device were accepted, normalized, semantically processed with local Phi-3 Mini first, validated, stored, translated, and exposed to the customer console. The complete safe load-management path and dataspace export also completed.",
    )
    add_body(
        document,
        "The final result was 30 of 30 telemetry messages accepted, 120 of 120 unique semantic and IEEE records stored, zero dropped messages, zero dropped readings, zero duplicate rows, zero processing errors, and 100% end-to-end completion. No real household device was controlled.",
    )
    add_callout(
        document,
        "Final assessment",
        "The architecture is ready for larger controlled local testing. Semantic inference concurrency and ontology coverage should be improved before substantially increasing the workload or moving toward production deployment.",
    )

    document.add_heading("Appendix A. Reproduction Commands", level=1)
    commands = [
        r"cd C:\Users\Mani\Desktop\Github\IOT-PIPELINE",
        r"powershell -ExecutionPolicy Bypass -File .\scripts\start-demo.ps1 -Build",
        r"powershell -ExecutionPolicy Bypass -File .\scripts\check-health.ps1",
        r"powershell -ExecutionPolicy Bypass -File .\scripts\run-multi-household-validation.ps1 -Households 5 -Cycles 2",
        r"cd .\apps\customer-console",
        r"npm run build",
        r"npm run start",
        r"npm run smoke:multi-household",
    ]
    for command in commands:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)

    document.add_heading("Appendix B. Evidence Locations", level=1)
    add_bullets(
        document,
        [
            "Machine-readable results: docs/demo-assets/multi-household-validation-results.json",
            "Dashboard screenshot: docs/demo-assets/multi-household-dashboard.png",
            "Technical report: docs/multi-household-scalability-validation-report.md",
            "Validation runner: scripts/run-multi-household-validation.js",
            "PowerShell wrapper: scripts/run-multi-household-validation.ps1",
        ],
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
